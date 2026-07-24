"""Export artifacts to Word (.docx) and PDF.

Design note worth stating, because it is the whole trick here:

`render.py` already turns every structured payload into markdown, and that markdown is what
the reviewer sees on screen and what the diff engine compares. Writing a *second* renderer
per format per artifact type would be nine types x three formats of drift waiting to happen —
the PDF would quietly disagree with the screen.

So instead: parse that one markdown into a small block model (heading / paragraph / bullets /
ordered / table), and render the block model into Word and PDF. One source of truth, three
outputs, guaranteed identical content. The markdown grammar is closed and machine-generated,
so parsing it is safe in a way that parsing arbitrary markdown would not be.
"""

from __future__ import annotations

import base64
import io
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Literal

from app.core.logging import log
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.models import ArtifactType, ArtifactVersion

NAVY = RGBColor(0x00, 0x4C, 0x8F)
GREY = RGBColor(0x6B, 0x7A, 0x90)
NAVY_HEX = "004C8F"

TITLES: dict[str, str] = {
    ArtifactType.BUSINESS_REQUIREMENTS.value: "Structured Business Requirements",
    ArtifactType.CONCEPT_NOTE.value: "Concept Note",
    ArtifactType.WIREFRAME.value: "Wireframe Specification",
    ArtifactType.BRD.value: "Business Requirements Document",
    ArtifactType.FRD.value: "Functional Requirements Document",
    ArtifactType.SRS.value: "Software Requirements Specification",
    ArtifactType.USER_STORIES.value: "User Stories",
    ArtifactType.ACCEPTANCE_CRITERIA.value: "Acceptance Criteria",
    ArtifactType.API_REQUIREMENTS.value: "API Requirements",
    ArtifactType.NFR.value: "Non-Functional Requirements",
    ArtifactType.SPRINT_PLAN.value: "Sprint Plan",
}

# The canonical order of a requirements pack — what a reviewer expects to page through.
PACK_ORDER = [
    ArtifactType.BUSINESS_REQUIREMENTS, ArtifactType.CONCEPT_NOTE, ArtifactType.WIREFRAME,
    ArtifactType.BRD, ArtifactType.FRD, ArtifactType.SRS, ArtifactType.USER_STORIES,
    ArtifactType.ACCEPTANCE_CRITERIA, ArtifactType.API_REQUIREMENTS, ArtifactType.NFR,
    ArtifactType.SPRINT_PLAN,
]


# ────────────────────────────── block model ───────────────────────────────────
@dataclass
class Block:
    kind: Literal["h1", "h2", "h3", "p", "ul", "ol", "table", "pagebreak"]
    text: str = ""
    items: list[str] = field(default_factory=list)
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


INLINE_RE = re.compile(r"\*\*(.+?)\*\*|_(.+?)_|`(.+?)`")

# ![alt](src) — the wireframe screenshots. These live INSIDE table cells, which is why they were
# invisible in both exports: the cell text was escaped, so the reader got the literal characters
# "![Dashboard](data:image/png;base64,iVBOR..." or, more often, nothing at all.
IMG_RE = re.compile(r"!\[([^\]]*)\]\((?P<src>data:image/[^)\s]+|https?://[^)\s]+)\)")


def image_src(cell: str) -> str | None:
    """The image source in this cell, if the cell IS an image."""
    m = IMG_RE.search(cell or "")
    return m.group("src") if m else None


_IMG_CACHE: dict[str, "bytes | None"] = {}


def image_bytes(src: str) -> bytes | None:
    """Raw bytes for an image, served from the prefetch cache when we have it.

    Word needs actual BYTES — it cannot follow a URL and it cannot read SVG. A remote URL is fetched
    at most once: prefetch_images() fills the cache concurrently before a build, and anything not
    already there is fetched on demand here.
    """
    if src in _IMG_CACHE:
        return _IMG_CACHE[src]
    return _image_bytes_uncached(src)


def _image_bytes_uncached(src: str) -> bytes | None:
    if src.startswith("data:"):
        head, _, b64 = src.partition(",")
        if "svg" in head:
            return None                    # python-docx cannot embed SVG. Do not pretend otherwise.
        try:
            return base64.b64decode(b64)
        except Exception:
            return None
    try:                                   # a real Stitch download URL
        import httpx

        r = httpx.get(src, timeout=6.0, follow_redirects=True)
        r.raise_for_status()
        return r.content
    except Exception as e:
        log.warning("export.image_fetch_failed", src=src[:60], error=str(e))
        return None


def prefetch_images(versions) -> None:
    """Fetch every remote image referenced by these documents concurrently, once, up front.

    The Word exporter embeds Stitch screenshots, which live at remote URLs that can be slow or dead.
    Fetched serially at 6s each, six screens is up to ~36s of dead wait — long enough that the
    browser's download fetch() gives up with a bare "Failed to fetch". A bounded pool collapses that
    to about one timeout's worth, and the results are cached so the builder itself never blocks.
    """
    import concurrent.futures

    urls: list[str] = []
    for v in versions:
        for b in parse(v.rendered_md or ""):
            cells = [b.text or ""] + [c for row in (b.rows or []) for c in row]
            for c in cells:
                u = image_src(c)
                if u and u.startswith(("http://", "https://")):
                    urls.append(u)
    todo = [u for u in dict.fromkeys(urls) if u not in _IMG_CACHE]
    if not todo:
        return
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as ex:
        futs = {ex.submit(_image_bytes_uncached, u): u for u in todo}
        for fut in concurrent.futures.as_completed(futs):
            try:
                _IMG_CACHE[futs[fut]] = fut.result()
            except Exception:
                _IMG_CACHE[futs[fut]] = None


def strip_inline(s: str) -> str:
    """Word and PDF get their emphasis from styles, not from asterisks.

    Image syntax is left intact: the exporters need to recognise it and embed the picture, and a
    base64 payload must never be run through the emphasis regex.
    """
    if IMG_RE.search(s or ""):
        return s
    return INLINE_RE.sub(lambda m: m.group(1) or m.group(2) or m.group(3), s).replace("<br>", "\n")


def parse(md: str) -> list[Block]:
    lines = md.split("\n")
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # table: | a | b |  followed by |---|---|
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            cells = lambda r: [strip_inline(c.strip()) for c in r.split("|")[1:-1]]  # noqa: E731
            headers = cells(line)
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(cells(lines[i]))
                i += 1
            blocks.append(Block("table", headers=headers, rows=rows))
            continue

        if m := re.match(r"^(#{1,3})\s+(.*)$", line):
            blocks.append(Block(f"h{len(m.group(1))}", text=strip_inline(m.group(2))))  # type: ignore[arg-type]
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(strip_inline(re.sub(r"^[-*]\s+", "", lines[i])))
                i += 1
            blocks.append(Block("ul", items=items))
            continue

        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(strip_inline(re.sub(r"^\d+\.\s+", "", lines[i])))
                i += 1
            blocks.append(Block("ol", items=items))
            continue

        blocks.append(Block("p", text=strip_inline(line)))
        i += 1
    return blocks


# ─────────────────────────────── Word (.docx) ─────────────────────────────────
def _shade(cell, hex_colour: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def _docx_blocks(doc: Document, blocks: list[Block]) -> None:
    for b in blocks:
        if b.kind == "pagebreak":
            doc.add_page_break()
        elif b.kind in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[b.kind]
            h = doc.add_heading(b.text, level=level)
            for run in h.runs:
                run.font.color.rgb = NAVY
        elif b.kind == "p":
            if (src := image_src(b.text)) and (raw := image_bytes(src)):
                doc.add_paragraph().add_run().add_picture(io.BytesIO(raw), width=Inches(5.4))
            else:
                doc.add_paragraph(b.text)
        elif b.kind == "ul":
            for it in b.items:
                doc.add_paragraph(it, style="List Bullet")
        elif b.kind == "ol":
            for it in b.items:
                doc.add_paragraph(it, style="List Number")
        elif b.kind == "table" and b.headers:
            t = doc.add_table(rows=1, cols=len(b.headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, htxt in enumerate(b.headers):
                cell = t.rows[0].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(htxt)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, NAVY_HEX)
            for row in b.rows:
                cells = t.add_row().cells
                for j, val in enumerate(row[: len(b.headers)]):
                    cells[j].text = ""
                    para = cells[j].paragraphs[0]
                    if (src := image_src(val)) and (raw := image_bytes(src)):
                        # The wireframe itself, in the cell. This is the whole point of the export.
                        para.add_run().add_picture(io.BytesIO(raw), width=Inches(4.2))
                        continue
                    r = para.add_run(val)
                    r.font.size = Pt(9)


def _docx_cover(doc: Document, *, title: str, project: str, meta: dict[str, Any]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HDFC BANK")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(project)
    r.font.size = Pt(14)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta.items():
        cells = t.add_row().cells
        rk = cells[0].paragraphs[0].add_run(f"{k}  ")
        rk.bold = True
        rk.font.size = Pt(9)
        rk.font.color.rgb = GREY
        rv = cells[1].paragraphs[0].add_run(str(v))
        rv.font.size = Pt(9)
    doc.add_page_break()


def _footer(doc: Document, text: str) -> None:
    p = doc.sections[0].footer.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(7.5)
        r.font.color.rgb = GREY


def to_docx(versions: list[ArtifactVersion], *, project_name: str, pack: bool = False) -> bytes:
    prefetch_images(versions)   # warm the image cache concurrently before we block on embeds
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    first = versions[0]
    title = "Requirements Pack" if pack else TITLES.get(first.artifact.type.value, first.artifact.type.value)
    meta = {
        "Project": project_name,
        "Generated": datetime.now(timezone.utc).strftime("%d %b %Y, %H:%M UTC"),
        "Status": "APPROVED" if all(v.approved for v in versions) else "DRAFT — PENDING APPROVAL",
    }
    if not pack:
        meta |= {"Version": f"v{first.version}", "Produced by": first.produced_by, "Model": first.model}
    else:
        meta |= {"Documents": str(len(versions))}

    _docx_cover(doc, title=title, project=project_name, meta=meta)
    _footer(doc, f"HDFC Bank · {project_name} · {title} · Generated by the Agentic SDLC Platform")

    for idx, v in enumerate(versions):
        if pack:
            h = doc.add_heading(TITLES.get(v.artifact.type.value, v.artifact.type.value), level=1)
            for r in h.runs:
                r.font.color.rgb = NAVY
            sub = doc.add_paragraph()
            r = sub.add_run(
                f"v{v.version} · {v.produced_by} · {v.model} · "
                f"{'APPROVED' if v.approved else 'PENDING APPROVAL'}"
            )
            r.font.size = Pt(8)
            r.font.color.rgb = GREY
        _docx_blocks(doc, parse(v.rendered_md))
        if pack and idx < len(versions) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────── PDF ─────────────────────────────────────
CSS = """
img.shot { width: 430px; border: 1px solid #DCE4EE; border-radius: 4px; }
img.shot.wide { width: 480px; }
@page { size: A4; margin: 20mm 16mm 18mm 16mm;
  @bottom-center { content: "HDFC Bank · Agentic SDLC Platform · Page " counter(page) " of " counter(pages);
                   font-size: 7.5pt; color: #6b7a90; font-family: Helvetica, Arial, sans-serif; } }
@page :first { @bottom-center { content: ""; } }
body { font-family: Helvetica, Arial, sans-serif; font-size: 9.5pt; color: #12263f; line-height: 1.5; }
.cover { height: 235mm; display: flex; flex-direction: column; justify-content: center; text-align: center; }
.cover .bank { color: #004C8F; font-weight: 700; letter-spacing: .18em; font-size: 11pt; }
.cover h1 { color: #004C8F; font-size: 30pt; margin: 10px 0 6px; }
.cover .proj { color: #6b7a90; font-size: 14pt; margin-bottom: 28px; }
.cover .meta { display: inline-block; text-align: left; font-size: 9pt; color: #12263f; }
.cover .meta div { padding: 3px 0; }
.cover .meta b { color: #6b7a90; display: inline-block; width: 110px; font-weight: 600; }
.cover .rule { width: 60px; height: 3px; background: #ED232A; margin: 0 auto 22px; }
h1 { color: #004C8F; font-size: 17pt; margin: 0 0 10px; page-break-after: avoid; }
h2 { color: #002E56; font-size: 12.5pt; margin: 18px 0 6px; page-break-after: avoid; }
h3 { color: #12263f; font-size: 10.5pt; margin: 13px 0 5px; page-break-after: avoid; }
.docmeta { font-size: 8pt; color: #6b7a90; margin-bottom: 14px; }
table { border-collapse: collapse; width: 100%; margin: 9px 0 14px; font-size: 8.5pt; }
th { background: #004C8F; color: #fff; text-align: left; padding: 6px 7px; font-weight: 600; }
td { border: 1px solid #dbe1e9; padding: 5px 7px; vertical-align: top; }
tr:nth-child(even) td { background: #f7f9fc; }
thead { display: table-header-group; }  /* repeat headers across page breaks */
tr { page-break-inside: avoid; }
ul, ol { margin: 5px 0 10px 16px; padding: 0; }
li { margin: 2px 0; }
.newdoc { page-break-before: always; }
"""


def _html_esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")


def _blocks_to_html(blocks: list[Block]) -> str:
    out = []
    for b in blocks:
        if b.kind in ("h1", "h2", "h3"):
            out.append(f"<{b.kind}>{_html_esc(b.text)}</{b.kind}>")
        elif b.kind == "p":
            if src := image_src(b.text):
                out.append(f'<p><img class="shot wide" src="{src}"/></p>')
            else:
                out.append(f"<p>{_html_esc(b.text)}</p>")
        elif b.kind in ("ul", "ol"):
            items = "".join(f"<li>{_html_esc(i)}</li>" for i in b.items)
            out.append(f"<{b.kind}>{items}</{b.kind}>")
        elif b.kind == "table" and b.headers:
            head = "".join(f"<th>{_html_esc(h)}</th>" for h in b.headers)
            body = "".join(
                "<tr>" + "".join(
                    (f'<td><img class="shot" src="{image_src(c)}"/></td>' if image_src(c)
                     else f"<td>{_html_esc(c)}</td>")
                    for c in r[: len(b.headers)]
                ) + "</tr>"
                for r in b.rows
            )
            out.append(f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>")
    return "\n".join(out)


def build_html(versions: list[ArtifactVersion], *, project_name: str, pack: bool = False) -> str:
    first = versions[0]
    title = "Requirements Pack" if pack else TITLES.get(first.artifact.type.value, first.artifact.type.value)
    approved = all(v.approved for v in versions)

    meta_rows = [
        ("Project", project_name),
        ("Generated", datetime.now(timezone.utc).strftime("%d %b %Y, %H:%M UTC")),
        ("Status", "APPROVED" if approved else "DRAFT — PENDING APPROVAL"),
    ]
    if pack:
        meta_rows.append(("Documents", str(len(versions))))
    else:
        meta_rows += [("Version", f"v{first.version}"), ("Produced by", first.produced_by),
                      ("Model", first.model)]

    meta_html = "".join(f"<div><b>{k}</b> {_html_esc(str(v))}</div>" for k, v in meta_rows)

    body = []
    for i, v in enumerate(versions):
        cls = "newdoc" if (pack and i > 0) else ""
        section = [f'<section class="{cls}">']
        if pack:
            section.append(f"<h1>{TITLES.get(v.artifact.type.value, v.artifact.type.value)}</h1>")
            section.append(
                f'<div class="docmeta">v{v.version} · {v.produced_by} · {v.model} · '
                f'{"APPROVED" if v.approved else "PENDING APPROVAL"}</div>'
            )
        section.append(_blocks_to_html(parse(v.rendered_md)))
        section.append("</section>")
        body.append("\n".join(section))

    return f"""<!doctype html><html><head><meta charset="utf-8"><style>{CSS}</style></head><body>
<div class="cover">
  <div class="bank">HDFC BANK</div>
  <h1>{_html_esc(title)}</h1>
  <div class="proj">{_html_esc(project_name)}</div>
  <div class="rule"></div>
  <div class="meta">{meta_html}</div>
</div>
{"".join(body)}
</body></html>"""


class PdfEngineUnavailable(RuntimeError):
    """WeasyPrint could not load — almost always its native libraries (Pango/Cairo/GDK-Pixbuf)
    missing from the host. The caller turns this into a clear message and offers Word instead."""


def to_pdf(versions: list[ArtifactVersion], *, project_name: str, pack: bool = False) -> bytes:
    from app.core.config import settings

    # Default to the low-memory engine: a beautiful PDF that OOM-kills the worker (bare "Internal
    # Server Error", uncatchable) is worse than a plain one that downloads every time.
    if settings.pdf_engine != "weasyprint":
        return _pdf_light(versions, project_name=project_name, pack=pack)
    try:
        from weasyprint import HTML  # heavy, and its native deps may be absent on a slim host
    except OSError as e:
        return _pdf_light(versions, project_name=project_name, pack=pack)

    html = build_html(versions, project_name=project_name, pack=pack)
    try:
        return HTML(string=html).write_pdf()
    except (OSError, MemoryError) as e:
        # A native-lib failure OR a memory failure: fall back to the low-memory engine rather than
        # denying the download. (A hard OOM kills the worker before we get here; this catches the
        # softer MemoryError and any lib fault.)
        from app.core.logging import log

        log.warning("pdf.weasyprint_fallback", error=str(e)[:160])
        return _pdf_light(versions, project_name=project_name, pack=pack)


def filename(v: ArtifactVersion, ext: str) -> str:
    return f"{v.artifact.type.value}_v{v.version}.{ext}"


# ── Low-memory PDF (fpdf2) ──────────────────────────────────────────────────────
def _pdf_light(versions: list["ArtifactVersion"], *, project_name: str, pack: bool) -> bytes:
    """A PDF built directly from the block model with fpdf2.

    WeasyPrint renders a full browser layout engine in memory; on a 512MB host a multi-document
    pack OOM-kills the worker, which surfaces as a bare "Internal Server Error" no application
    handler can catch — the process is already dead. fpdf2 is pure Python and streams text, so it
    fits. Plainer than WeasyPrint, which is the right trade: a plain PDF that downloads beats a
    beautiful one that kills the server.

    Tables are rendered as text lines, not fixed-width cells: fpdf2 raises "not enough horizontal
    space" the moment a column is narrower than a single word, and requirement tables have long
    cells. Reliability over grid lines.
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    NAVY = (0, 76, 143)
    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.set_margins(18, 16, 18)
    EPW = pdf.epw  # effective page width fpdf itself computes from the margins

    # Map the unicode punctuation the mock/LLM emits to latin-1 so it renders as the real glyph
    # instead of "?" — the core PDF fonts are latin-1 only.
    _SUBS = {
        "\u2014": "-", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u2022": "-",
        "\u00a0": " ", "\u2192": "->", "\u2265": ">=", "\u2264": "<=",
    }

    def _t(s: str) -> str:
        s = s or ""
        for k, v in _SUBS.items():
            s = s.replace(k, v)
        s = s.encode("latin-1", "replace").decode("latin-1")
        # Hard-break any token too long to fit, so multi_cell never fails on one word (e.g. a URL).
        return re.sub(r"(\S{55})(?=\S)", r"\1 ", s)

    def _para(txt: str, *, size=10, style="", color=(20, 20, 20), h=5):
        pdf.set_font("Helvetica", style, size)
        pdf.set_text_color(*color)
        # new_x=LMARGIN, new_y=NEXT is essential: without it fpdf2 leaves the cursor at the right
        # edge of the last line, and the next paragraph starts in the right margin and overflows.
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(EPW, h, _t(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    for v in versions:
        pdf.add_page()
        _para("HDFC BANK", size=9, style="B", color=NAVY, h=5)
        title = v.artifact.type.value.replace("_", " ").title()
        _para(f"{title}  -  {project_name}", size=15, style="B", color=NAVY, h=7)
        _para(f"v{v.version} · {v.produced_by} · {v.model}", size=8, color=(90, 90, 90), h=4)
        pdf.ln(1)
        pdf.set_draw_color(210, 210, 210)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)

        for b in parse(v.rendered_md or ""):
            if b.kind in ("h1", "h2", "h3"):
                pdf.ln(1.5)
                _para(b.text, size={"h1": 14, "h2": 12, "h3": 11}[b.kind], style="B", color=NAVY, h=6)
            elif b.kind == "p":
                _para(b.text, h=5)
                pdf.ln(0.5)
            elif b.kind in ("ul", "ol"):
                for j, it in enumerate(b.items):
                    _para((f"{j+1}. " if b.kind == "ol" else "-  ") + it, h=5)
                pdf.ln(0.5)
            elif b.kind == "table" and b.headers:
                _para(" | ".join(b.headers), size=8, style="B", color=NAVY, h=5)
                for row in b.rows:
                    _para(" | ".join(str(c) for c in row), size=8, color=(60, 60, 60), h=4.5)
                pdf.ln(1)

    return bytes(pdf.output())
